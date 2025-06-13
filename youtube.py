from docx import Document
import requests
from config import API_KEY
 
PLAYLIST_ID = 'PLPe9IkX86X3y5m_MvtNu2ughxsvkqUNKr'  # Your playlist ID
URL = 'https://www.googleapis.com/youtube/v3/playlistItems'

titles = []
next_page_token = ''

print("Fetching playlist videos...")

while True:
    params = {
        'part': 'snippet',
        'playlistId': PLAYLIST_ID,
        'maxResults': 50,
        'pageToken': next_page_token,
        'key': API_KEY
    }

    response = requests.get(URL, params=params)
    data = response.json()

    for item in data['items']:
        title = item['snippet']['title']
        titles.append(title)

    next_page_token = data.get('nextPageToken')
    if not next_page_token:
        break

# Save as a Word Document
doc = Document()
doc.add_heading('YouTube Playlist Video Titles', 0)

for idx, title in enumerate(titles, 1):
    doc.add_paragraph(f"{idx}. {title}")

filename = "playlist_titles.docx"
doc.save(filename)
print(f"\n✅ Done! Titles saved in: {filename}")



